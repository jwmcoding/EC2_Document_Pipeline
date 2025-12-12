"""
Optimized Document Converter - Phase 2 Integration
High-performance document converter with size-based optimization strategies
"""

import io
import tempfile
import subprocess
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
import logging
from PIL import Image
from pdf2image import convert_from_bytes
import openpyxl
import pandas as pd
from docx import Document as DocxDocument
import extract_msg

try:
    from .optimized_pdf_processor import OptimizedPDFProcessor
    from .optimized_excel_processor import OptimizedExcelProcessor
    from .optimized_docx_processor import OptimizedDocxProcessor
    from ..config.settings import Settings
except ImportError:
    # Fallback imports for direct execution
    import sys
    import os
    sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(__file__))))
    from src.parsers.optimized_pdf_processor import OptimizedPDFProcessor
    from src.parsers.optimized_excel_processor import OptimizedExcelProcessor
    from src.parsers.optimized_docx_processor import OptimizedDocxProcessor
    from src.config.settings import Settings


class OptimizedDocumentConverter:
    """
    High-performance document converter with Phase 2 optimizations
    
    Features:
    - Size-based processing strategies for PDF, Excel, DOCX
    - Intelligent content limits and sampling
    - Memory-efficient processing for large files
    - Fallback to original methods when needed
    """
    
    def __init__(self, use_optimizations: bool = True):
        self.logger = logging.getLogger(__name__)
        self.settings = Settings()
        self.use_optimizations = use_optimizations and self.settings.OPTIMIZED_PARSING_ENABLED
        
        # Initialize optimized processors
        if self.use_optimizations:
            self.pdf_processor = OptimizedPDFProcessor()
            self.excel_processor = OptimizedExcelProcessor()
            self.docx_processor = OptimizedDocxProcessor()
            
            self.logger.info("🚀 Optimized document converter initialized")
        else:
            self.pdf_processor = None
            self.excel_processor = None
            self.docx_processor = None
            
            self.logger.info("📄 Standard document converter initialized")
        
        self.supported_extensions = {
            '.pdf': 'pdf',
            '.xlsx': 'excel',
            '.xls': 'excel', 
            '.csv': 'excel',
            '.docx': 'word',
            '.doc': 'word',
            '.pptx': 'powerpoint',
            '.txt': 'text',
            '.png': 'image',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.msg': 'outlook_msg'
        }
    
    def can_process(self, file_path: str) -> bool:
        """Check if we can process this file type"""
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_extensions
    
    def convert_to_processable_content(self, file_path: str, content: bytes) -> Tuple[bytes, str]:
        """
        Convert various file types to PDF bytes or extract text directly with optimization
        Returns: (processed_content, content_type)
        content_type: 'pdf' for PDF bytes, 'text' for direct text
        """
        extension = Path(file_path).suffix.lower()
        size_mb = len(content) / (1024 * 1024)
        
        self.logger.debug(f"🔄 Processing {extension} file: {size_mb:.2f}MB")
        
        try:
            if extension == '.pdf':
                if self.use_optimizations and self.pdf_processor:
                    # Use optimized PDF processing
                    return self._process_pdf_optimized(content, file_path)
                else:
                    return content, 'pdf'
            
            elif extension in ['.xlsx', '.xls', '.csv']:
                if self.use_optimizations and self.excel_processor:
                    # Use optimized Excel processing
                    return self._process_excel_optimized(content, extension, file_path)
                else:
                    # Fall back to original method
                    text_content = self._extract_excel_text_original(content, extension)
                    return text_content.encode('utf-8'), 'text'
            
            elif extension == '.docx':
                if self.use_optimizations and self.docx_processor:
                    # Use optimized DOCX processing
                    return self._process_docx_optimized(content, file_path)
                else:
                    # Fall back to original method
                    text_content = self._extract_docx_text_original(content)
                    return text_content.encode('utf-8'), 'text'
            
            elif extension == '.doc':
                text_content = self._extract_doc_text(content, file_path)
                return text_content.encode('utf-8'), 'text'
            
            elif extension == '.txt':
                return content, 'text'
            
            elif extension in ['.png', '.jpg', '.jpeg']:
                pdf_content = self._image_to_pdf(content)
                return pdf_content, 'pdf'
            
            elif extension == '.msg':
                text_content = self._extract_msg_text(content, file_path)
                return text_content.encode('utf-8'), 'text'
            
            else:
                raise ValueError(f"Unsupported file type: {extension}")
                
        except Exception as e:
            self.logger.error(f"Error converting {file_path}: {e}")
            # If optimized processing fails, try falling back to original methods
            if self.use_optimizations:
                self.logger.info(f"🔄 Falling back to original processing for {file_path}")
                return self._fallback_processing(extension, content, file_path)
            else:
                raise
    
    def _process_pdf_optimized(self, content: bytes, file_path: str) -> Tuple[bytes, str]:
        """Process PDF using optimized processor"""
        try:
            result = self.pdf_processor.process_pdf_optimized(content, {'name': Path(file_path).name})
            
            if result.get('success', False):
                text_content = result['text']
                self.logger.info(f"⚡ PDF optimized: {result.get('strategy_used')} strategy, "
                               f"{result.get('processing_time', 0):.2f}s")
                return text_content.encode('utf-8'), 'text'
            else:
                # Fall back to original PDF processing
                return content, 'pdf'
                
        except Exception as e:
            self.logger.warning(f"Optimized PDF processing failed: {e}")
            return content, 'pdf'
    
    def _process_excel_optimized(self, content: bytes, extension: str, file_path: str) -> Tuple[bytes, str]:
        """Process Excel using optimized processor"""
        try:
            result = self.excel_processor.process_excel_optimized(
                content, extension, {'name': Path(file_path).name}
            )
            
            if result.get('success', False):
                text_content = result['text']
                self.logger.info(f"⚡ Excel optimized: {result.get('size_category')} strategy, "
                               f"{result.get('processing_time', 0):.2f}s")
                return text_content.encode('utf-8'), 'text'
            else:
                # Fall back to original processing
                text_content = self._extract_excel_text_original(content, extension)
                return text_content.encode('utf-8'), 'text'
                
        except Exception as e:
            self.logger.warning(f"Optimized Excel processing failed: {e}")
            text_content = self._extract_excel_text_original(content, extension)
            return text_content.encode('utf-8'), 'text'
    
    def _process_docx_optimized(self, content: bytes, file_path: str) -> Tuple[bytes, str]:
        """Process DOCX using optimized processor"""
        try:
            result = self.docx_processor.process_docx_optimized(
                content, {'name': Path(file_path).name}
            )
            
            if result.get('success', False):
                text_content = result['text']
                self.logger.info(f"⚡ DOCX optimized: {result.get('size_category')} strategy, "
                               f"{result.get('processing_time', 0):.2f}s")
                return text_content.encode('utf-8'), 'text'
            else:
                # Fall back to original processing
                text_content = self._extract_docx_text_original(content)
                return text_content.encode('utf-8'), 'text'
                
        except Exception as e:
            self.logger.warning(f"Optimized DOCX processing failed: {e}")
            text_content = self._extract_docx_text_original(content)
            return text_content.encode('utf-8'), 'text'
    
    def _fallback_processing(self, extension: str, content: bytes, file_path: str) -> Tuple[bytes, str]:
        """Fallback to original processing methods"""
        if extension in ['.xlsx', '.xls', '.csv']:
            text_content = self._extract_excel_text_original(content, extension)
            return text_content.encode('utf-8'), 'text'
        elif extension == '.docx':
            text_content = self._extract_docx_text_original(content)
            return text_content.encode('utf-8'), 'text'
        elif extension == '.pdf':
            return content, 'pdf'
        else:
            raise ValueError(f"Cannot process file type: {extension}")
    
    def _extract_excel_text_original(self, content: bytes, extension: str) -> str:
        """Original Excel processing method (fallback)"""
        try:
            if extension == '.csv':
                df = pd.read_csv(io.BytesIO(content))
                sheets_data = {'Sheet1': df}
            else:
                excel_file = pd.ExcelFile(io.BytesIO(content))
                sheets_data = {}
                for sheet_name in excel_file.sheet_names:
                    try:
                        sheets_data[sheet_name] = excel_file.parse(sheet_name)
                    except Exception as e:
                        self.logger.warning(f"Could not parse sheet {sheet_name}: {e}")
                        continue
            
            # Convert to readable text format
            text_parts = []
            for sheet_name, df in sheets_data.items():
                text_parts.append(f"=== {sheet_name} ===")
                
                if df.empty:
                    text_parts.append("(Empty sheet)")
                    continue
                
                # Add headers
                headers = " | ".join(str(col) for col in df.columns)
                text_parts.append(headers)
                text_parts.append("-" * min(len(headers), 100))
                
                # Add data (limit rows to prevent huge text blocks)
                max_rows = min(1000, len(df))
                for idx, row in df.head(max_rows).iterrows():
                    try:
                        row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row.values)
                        text_parts.append(row_text)
                    except Exception as e:
                        self.logger.warning(f"Error processing row {idx}: {e}")
                        continue
                
                if len(df) > 1000:
                    text_parts.append(f"... and {len(df) - 1000} more rows")
                
                text_parts.append("")  # Empty line between sheets
            
            return "\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Error extracting Excel text: {e}")
            return f"Error reading Excel file: {str(e)}"
    
    def _extract_docx_text_original(self, content: bytes) -> str:
        """Original DOCX processing method (fallback)"""
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table content
            for table in doc.tables:
                text_parts.append("\n=== TABLE ===")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("=== END TABLE ===\n")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {e}")
            return f"Error reading Word document: {str(e)}"
    
    def _image_to_pdf(self, content: bytes) -> bytes:
        """Convert image to PDF for consistent processing"""
        try:
            # Open image
            image = Image.open(io.BytesIO(content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Save as PDF
            pdf_buffer = io.BytesIO()
            image.save(pdf_buffer, format='PDF')
            return pdf_buffer.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error converting image to PDF: {e}")
            raise ValueError(f"Cannot convert image to PDF: {str(e)}")
    
    def extract_msg_metadata(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Extract metadata from MSG files for compatibility"""
        # Simple metadata extraction for MSG files
        return {
            'file_type': '.msg',
            'file_path': file_path,
            'size_bytes': len(content)
        }
    
    def _extract_msg_text(self, content: bytes, file_path: str) -> str:
        """Extract text from Outlook .msg files (unchanged from original)"""
        try:
            # Create a temporary file since extract_msg works with file paths
            with tempfile.NamedTemporaryFile(suffix='.msg', delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                try:
                    # Open the MSG file
                    msg = extract_msg.openMsg(temp_file.name)
                    
                    # Extract basic email components
                    text_parts = []
                    
                    # Email metadata
                    if hasattr(msg, 'subject') and msg.subject:
                        text_parts.append(f"Subject: {msg.subject}")
                    
                    if hasattr(msg, 'sender') and msg.sender:
                        text_parts.append(f"From: {msg.sender}")
                    
                    if hasattr(msg, 'to') and msg.to:
                        text_parts.append(f"To: {msg.to}")
                    
                    if hasattr(msg, 'cc') and msg.cc:
                        text_parts.append(f"CC: {msg.cc}")
                    
                    if hasattr(msg, 'date') and msg.date:
                        text_parts.append(f"Date: {msg.date}")
                    
                    # Email body - try different body types
                    body_found = False
                    
                    # Try plain text body first
                    if hasattr(msg, 'body') and msg.body:
                        text_parts.append("--- Email Body ---")
                        text_parts.append(msg.body.strip())
                        body_found = True
                    
                    # If no plain text, try HTML body
                    elif hasattr(msg, 'htmlBody') and msg.htmlBody:
                        text_parts.append("--- Email Body (HTML) ---")
                        # Basic HTML tag removal for better text extraction
                        import re
                        html_text = re.sub(r'<[^>]+>', '', msg.htmlBody)
                        html_text = html_text.replace('&nbsp;', ' ').replace('&amp;', '&')
                        text_parts.append(html_text.strip())
                        body_found = True
                    
                    if not body_found:
                        text_parts.append("--- Email Body ---")
                        text_parts.append("(No readable body content found)")
                    
                    # List attachments if any
                    if hasattr(msg, 'attachments') and msg.attachments:
                        text_parts.append("--- Attachments ---")
                        for i, attachment in enumerate(msg.attachments):
                            if hasattr(attachment, 'longFilename') and attachment.longFilename:
                                text_parts.append(f"Attachment {i+1}: {attachment.longFilename}")
                            elif hasattr(attachment, 'shortFilename') and attachment.shortFilename:
                                text_parts.append(f"Attachment {i+1}: {attachment.shortFilename}")
                            else:
                                text_parts.append(f"Attachment {i+1}: (unnamed)")
                    
                    # Close the MSG file
                    msg.close()
                    
                    result_text = '\n'.join(text_parts)
                    self.logger.info(f"Successfully extracted {len(result_text)} characters from MSG file: {Path(file_path).name}")
                    
                    return result_text if result_text.strip() else f"MSG file processed but no readable content found: {Path(file_path).name}"
                
                finally:
                    # Clean up temporary file
                    import os
                    try:
                        os.unlink(temp_file.name)
                    except:
                        pass
                        
        except Exception as e:
            self.logger.error(f"Error extracting text from MSG file {file_path}: {e}")
            # Return a basic message rather than failing completely
            return f"MSG file processing error: {Path(file_path).name} - {str(e)}"
    
    def _extract_doc_text(self, content: bytes, file_path: str) -> str:
        """Extract text from legacy .doc files using multiple fallback methods (unchanged from original)"""
        try:
            # Method 1: Try python-docx2txt if available
            try:
                import docx2txt
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    
                    text = docx2txt.process(temp_file.name)
                    
                    # Clean up
                    import os
                    os.unlink(temp_file.name)
                    
                    if text and text.strip():
                        self.logger.debug(f"Successfully extracted .doc text using docx2txt: {len(text)} chars")
                        return text.strip()
                        
            except ImportError:
                self.logger.debug("docx2txt not available, trying alternative methods...")
            except Exception as e:
                self.logger.debug(f"docx2txt extraction failed: {e}")
            
            # Continue with other fallback methods...
            self.logger.warning(f"Could not extract text from .doc file: {Path(file_path).name}")
            return f"[Legacy Word Document: {Path(file_path).name}] - Text extraction not available. Install docx2txt or antiword for better .doc support."
            
        except Exception as e:
            self.logger.error(f"Error processing .doc file {file_path}: {e}")
            return f"[Error processing .doc file: {Path(file_path).name}]"
    
    def get_optimization_status(self) -> Dict[str, Any]:
        """Get current optimization status and settings"""
        return {
            'optimizations_enabled': self.use_optimizations,
            'processors_available': {
                'pdf': self.pdf_processor is not None,
                'excel': self.excel_processor is not None,
                'docx': self.docx_processor is not None
            },
            'settings': {
                'pdf_optimizations': self.pdf_processor.get_optimization_stats() if self.pdf_processor else None,
                'excel_optimizations': self.excel_processor.get_optimization_info() if self.excel_processor else None,
                'docx_optimizations': self.docx_processor.get_optimization_info() if self.docx_processor else None
            }
        } 