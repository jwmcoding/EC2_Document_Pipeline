"""
Optimized DOCX Processor - Phase 2 Performance Enhancement
Fast DOCX processing with intelligent paragraph and table limits

Updated Dec 2025: Tables now formatted using unified table_formatter
for consistent detection and preservation during chunking.
"""

from docx import Document as DocxDocument
import io
import logging
from typing import Dict, Any, List, Optional
import time

from src.parsers.table_formatter import format_table_for_chunking, convert_docx_table_to_list


class OptimizedDocxProcessor:
    """
    High-performance DOCX processor with intelligent content limits
    
    Features:
    - Dynamic paragraph limits based on file size
    - Efficient table extraction with limits
    - Memory-efficient text processing
    - Smart content sampling for very large documents
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Performance thresholds (MB)
        self.size_thresholds = {
            'small': 2,    # < 2MB
            'medium': 10,  # 2-10MB
            'large': 30    # > 10MB
        }
        
        # NO CONTENT LIMITS - Process all data
        self.limits = {
            'small': {
                'paragraphs': None,  # NO LIMITS
                'tables': None,      # NO LIMITS
                'table_rows': None   # NO LIMITS
            },
            'medium': {
                'paragraphs': None,  # NO LIMITS
                'tables': None,      # NO LIMITS
                'table_rows': None   # NO LIMITS
            },
            'large': {
                'paragraphs': None,  # NO LIMITS
                'tables': None,      # NO LIMITS
                'table_rows': None   # NO LIMITS
            }
        }
    
    def process_docx_optimized(self, content: bytes, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """
        Process DOCX files with size-based optimization
        
        Args:
            content: File content as bytes
            metadata: Document metadata
            
        Returns:
            Processing result with text and metadata
        """
        start_time = time.time()
        size_mb = len(content) / (1024 * 1024)
        
        # Categorize file size
        category = self._categorize_file_size(size_mb)
        limits = self.limits[category]
        
        self.logger.info(f"📊 DOCX Strategy: {category} "
                        f"(size: {size_mb:.2f}MB, "
                        f"paragraph_limit: {limits['paragraphs']}, "
                        f"table_limit: {limits['tables']})")
        
        try:
            result = self._process_docx_with_limits(content, limits, metadata)
            
            processing_time = time.time() - start_time
            result['processing_time'] = processing_time
            result['size_category'] = category
            result['size_mb'] = size_mb
            
            self.logger.info(f"⚡ DOCX processed in {processing_time:.2f}s "
                           f"({category} strategy)")
            
            return result
            
        except Exception as e:
            self.logger.error(f"❌ DOCX processing failed: {e}")
            return self._create_error_result(str(e), metadata)
    
    def _categorize_file_size(self, size_mb: float) -> str:
        """Categorize file size for processing strategy"""
        if size_mb < self.size_thresholds['small']:
            return 'small'
        elif size_mb < self.size_thresholds['medium']:
            return 'medium'
        else:
            return 'large'
    
    def _process_docx_with_limits(self, content: bytes, limits: Dict[str, int],
                                metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Process DOCX with specified limits"""
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            
            # Process paragraphs with limit
            paragraphs_processed = self._extract_paragraphs_optimized(
                doc, limits['paragraphs'], text_parts
            )
            
            # Process tables with limit
            tables_processed = self._extract_tables_optimized(
                doc, limits['tables'], limits['table_rows'], text_parts
            )
            
            # Add processing summary for completed documents
            text_parts.append("\n=== PROCESSING SUMMARY ===")
            text_parts.append(f"Processed {paragraphs_processed}/{len(doc.paragraphs)} paragraphs")
            text_parts.append(f"Processed {tables_processed}/{len(doc.tables)} tables")
            text_parts.append("(Complete document processed - no truncation)")
            
            full_text = '\n'.join(text_parts)
            
            return {
                'text': full_text,
                'metadata': {
                    **metadata,
                    'parser': 'optimized_docx',
                    'total_paragraphs': len(doc.paragraphs),
                    'processed_paragraphs': paragraphs_processed,
                    'total_tables': len(doc.tables),
                    'processed_tables': tables_processed,
                    'truncated': False,  # No truncation applied
                    'text_length': len(full_text)
                },
                'success': True
            }
            
        except Exception as e:
            self.logger.error(f"DOCX processing error: {e}")
            return self._create_error_result(f"DOCX processing error: {str(e)}", metadata)
    
    def _extract_paragraphs_optimized(self, doc: DocxDocument, max_paragraphs: int,
                                    text_parts: List[str]) -> int:
        """Extract ALL paragraphs - NO LIMITS"""
        processed = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            # NO LIMIT CHECK - process all paragraphs
                
            paragraph_text = paragraph.text.strip()
            if paragraph_text:
                # NO character limit - preserve full paragraph content
                text_parts.append(paragraph_text)
                processed += 1
        
        return processed
    
    def _extract_tables_optimized(self, doc: DocxDocument, max_tables: int,
                                max_rows_per_table: int, text_parts: List[str]) -> int:
        """Extract ALL tables using unified formatter for chunker preservation"""
        processed = 0
        
        for i, table in enumerate(doc.tables):
            # NO TABLE LIMIT - process all tables
            
            try:
                # Convert DOCX table to 2D list
                table_data = convert_docx_table_to_list(table)
                
                # Format using unified table formatter for chunker detection
                formatted_table = format_table_for_chunking(
                    table_data,
                    f"TABLE_{i + 1}"
                )
                
                if formatted_table:
                    text_parts.append("")  # Blank line before table
                    text_parts.append(formatted_table)
                    self.logger.debug(f"Formatted TABLE_{i + 1} with {len(table_data)} rows")
                
                processed += 1
                
            except Exception as e:
                self.logger.warning(f"Error processing table {i+1}: {e}")
                text_parts.append(f"\n=== TABLE {i+1} ===")
                text_parts.append(f"Error reading table: {str(e)}")
                text_parts.append("=== END TABLE ===\n")
                processed += 1
        
        return processed
    
    def _create_error_result(self, error_msg: str, metadata: Dict[str, Any]) -> Dict[str, Any]:
        """Create standardized error result"""
        return {
            'text': f"Error processing DOCX file: {error_msg}",
            'metadata': {
                **metadata,
                'parser': 'optimized_docx_error',
                'error': error_msg
            },
            'success': False,
            'error': error_msg
        }
    
    def get_optimization_info(self) -> Dict[str, Any]:
        """Get information about optimization settings"""
        return {
            'size_thresholds_mb': self.size_thresholds,
            'content_limits': self.limits,
            'features': [
                'Dynamic paragraph limits based on file size',
                'Table extraction with row limits',
                'Cell content truncation for large cells',
                'Memory-efficient text processing',
                'Smart content sampling for large documents'
            ]
        } 