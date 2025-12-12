"""
Document Converter for Multiple File Types
Converts various file types to PDF or extracts text for PDFPlumber processing
"""

import io
import tempfile
import subprocess
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
import logging
from PIL import Image
from pdf2image import convert_from_bytes
import openpyxl
import pandas as pd
from docx import Document as DocxDocument
import extract_msg
from .powerpoint_parser import PowerPointParser
from .enhanced_powerpoint_parser import EnhancedPowerPointParser


class DocumentConverter:
    """Convert various file types to PDF or extract text for PDFPlumber processing"""
    
    def __init__(self, openai_client=None, enable_vision_analysis: bool = False, vision_model: str = "gpt-4o"):
        self.logger = logging.getLogger(__name__)
        self.powerpoint_parser = PowerPointParser()
        
        # Enhanced PowerPoint parser with VLM capabilities
        if openai_client and enable_vision_analysis:
            self.enhanced_powerpoint_parser = EnhancedPowerPointParser(
                openai_client=openai_client, 
                enable_vision_analysis=True,
                vision_model=vision_model
            )
        else:
            self.enhanced_powerpoint_parser = None
        self.supported_extensions = {
            '.pdf': 'native',
            '.xlsx': 'excel',
            '.xls': 'excel', 
            '.csv': 'excel',
            '.docx': 'word',
            '.doc': 'word',
            '.pptx': 'powerpoint',
            '.txt': 'text',
            '.png': 'image',
            '.jpg': 'image',
            '.jpeg': 'image',
            '.msg': 'outlook_msg'
        }
    
    def can_process(self, file_path: str, file_name: str = None) -> bool:
        """Check if we can process this file type
        
        Args:
            file_path: Path to the file (may not have extension for Salesforce exports)
            file_name: Original filename with extension (used when file_path lacks extension)
        """
        extension = Path(file_path).suffix.lower()
        if not extension and file_name:
            extension = Path(file_name).suffix.lower()
        return extension in self.supported_extensions
    
    def convert_to_processable_content(self, file_path: str, content: bytes, file_name: str = None) -> Tuple[bytes, str]:
        """
        Convert various file types to PDF bytes or extract text directly
        Returns: (processed_content, content_type)
        content_type: 'pdf' for PDF bytes, 'text' for direct text
        
        Args:
            file_path: Path to the file (may not have extension for Salesforce exports)
            content: File content bytes
            file_name: Original filename with extension (used when file_path lacks extension)
        """
        # First try extension from file_path, then from file_name
        extension = Path(file_path).suffix.lower()
        if not extension and file_name:
            extension = Path(file_name).suffix.lower()
        
        try:
            if extension == '.pdf':
                return content, 'pdf'
            
            elif extension in ['.xlsx', '.xls', '.csv']:
                text_content = self._extract_excel_text(content, extension)
                return text_content.encode('utf-8'), 'text'
            
            elif extension == '.docx':
                text_content = self._extract_docx_text(content)
                return text_content.encode('utf-8'), 'text'
            
            elif extension == '.doc':
                text_content = self._extract_doc_text(content, file_path)
                return text_content.encode('utf-8'), 'text'
            
            elif extension == '.txt':
                return content, 'text'
            
            elif extension in ['.png', '.jpg', '.jpeg']:
                # For now, convert images to PDF for consistent processing
                pdf_content = self._image_to_pdf(content)
                return pdf_content, 'pdf'
            
            elif extension == '.msg':
                text_content = self._extract_msg_text(content, file_path)
                return text_content.encode('utf-8'), 'text'
            
            elif extension == '.pptx':
                text_content = self._extract_pptx_text(content, file_path)
                return text_content.encode('utf-8'), 'text'
            
            else:
                raise ValueError(f"Unsupported file type: {extension}")
                
        except Exception as e:
            self.logger.error(f"Error converting {file_path}: {e}")
            raise
    
    def _extract_excel_text(self, content: bytes, extension: str) -> str:
        """Extract text from Excel/CSV files"""
        try:
            if extension == '.csv':
                df = pd.read_csv(io.BytesIO(content))
                sheets_data = {'Sheet1': df}
            else:
                excel_file = pd.ExcelFile(io.BytesIO(content))
                sheets_data = {}
                for sheet_name in excel_file.sheet_names:
                    try:
                        sheets_data[sheet_name] = excel_file.parse(sheet_name)
                    except Exception as e:
                        self.logger.warning(f"Could not parse sheet {sheet_name}: {e}")
                        continue
            
            # Convert to readable text format
            text_parts = []
            for sheet_name, df in sheets_data.items():
                text_parts.append(f"=== {sheet_name} ===")
                
                if df.empty:
                    text_parts.append("(Empty sheet)")
                    continue
                
                # Add headers
                headers = " | ".join(str(col) for col in df.columns)
                text_parts.append(headers)
                text_parts.append("-" * min(len(headers), 100))
                
                # Add data (limit rows to prevent huge text blocks)
                max_rows = min(1000, len(df))
                for idx, row in df.head(max_rows).iterrows():
                    try:
                        row_text = " | ".join(str(val) if pd.notna(val) else "" for val in row.values)
                        text_parts.append(row_text)
                    except Exception as e:
                        self.logger.warning(f"Error processing row {idx}: {e}")
                        continue
                
                if len(df) > 1000:
                    text_parts.append(f"... and {len(df) - 1000} more rows")
                
                text_parts.append("")  # Empty line between sheets
            
            return "\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Error extracting Excel text: {e}")
            return f"Error reading Excel file: {str(e)}"
    
    def _extract_docx_text(self, content: bytes) -> str:
        """Extract text from DOCX files"""
        try:
            doc = DocxDocument(io.BytesIO(content))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract table content
            for table in doc.tables:
                text_parts.append("\n=== TABLE ===")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("=== END TABLE ===\n")
            
            return "\n".join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX text: {e}")
            return f"Error reading Word document: {str(e)}"
    
    def _image_to_pdf(self, content: bytes) -> bytes:
        """Convert image to PDF for consistent processing"""
        try:
            # Open image
            image = Image.open(io.BytesIO(content))
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Save as PDF
            pdf_buffer = io.BytesIO()
            image.save(pdf_buffer, format='PDF')
            return pdf_buffer.getvalue()
            
        except Exception as e:
            self.logger.error(f"Error converting image to PDF: {e}")
            raise ValueError(f"Cannot convert image to PDF: {str(e)}")
    
    def _extract_msg_text(self, content: bytes, file_path: str) -> str:
        """Extract text from Outlook .msg files"""
        try:
            # Create a temporary file since extract_msg works with file paths
            with tempfile.NamedTemporaryFile(suffix='.msg', delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                try:
                    # Open the MSG file
                    msg = extract_msg.openMsg(temp_file.name)
                    
                    # Extract basic email components
                    text_parts = []
                    
                    # Email metadata
                    if hasattr(msg, 'subject') and msg.subject:
                        text_parts.append(f"Subject: {msg.subject}")
                    
                    if hasattr(msg, 'sender') and msg.sender:
                        text_parts.append(f"From: {msg.sender}")
                    
                    if hasattr(msg, 'to') and msg.to:
                        text_parts.append(f"To: {msg.to}")
                    
                    if hasattr(msg, 'cc') and msg.cc:
                        text_parts.append(f"CC: {msg.cc}")
                    
                    if hasattr(msg, 'date') and msg.date:
                        text_parts.append(f"Date: {msg.date}")
                    
                    # Email body - try different body types
                    body_found = False
                    
                    # Try plain text body first
                    if hasattr(msg, 'body') and msg.body:
                        text_parts.append("--- Email Body ---")
                        text_parts.append(msg.body.strip())
                        body_found = True
                    
                    # If no plain text, try HTML body
                    elif hasattr(msg, 'htmlBody') and msg.htmlBody:
                        text_parts.append("--- Email Body (HTML) ---")
                        # Basic HTML tag removal for better text extraction
                        import re
                        html_text = re.sub(r'<[^>]+>', '', msg.htmlBody)
                        html_text = html_text.replace('&nbsp;', ' ').replace('&amp;', '&')
                        text_parts.append(html_text.strip())
                        body_found = True
                    
                    if not body_found:
                        text_parts.append("--- Email Body ---")
                        text_parts.append("(No readable body content found)")
                    
                    # List attachments if any
                    if hasattr(msg, 'attachments') and msg.attachments:
                        text_parts.append("--- Attachments ---")
                        for i, attachment in enumerate(msg.attachments):
                            if hasattr(attachment, 'longFilename') and attachment.longFilename:
                                text_parts.append(f"Attachment {i+1}: {attachment.longFilename}")
                            elif hasattr(attachment, 'shortFilename') and attachment.shortFilename:
                                text_parts.append(f"Attachment {i+1}: {attachment.shortFilename}")
                            else:
                                text_parts.append(f"Attachment {i+1}: (unnamed)")
                    
                    # Close the MSG file
                    msg.close()
                    
                    result_text = '\n'.join(text_parts)
                    self.logger.info(f"Successfully extracted {len(result_text)} characters from MSG file: {Path(file_path).name}")
                    
                    return result_text if result_text.strip() else f"MSG file processed but no readable content found: {Path(file_path).name}"
                
                finally:
                    # Clean up temporary file
                    import os
                    try:
                        os.unlink(temp_file.name)
                    except:
                        pass
                        
        except Exception as e:
            self.logger.error(f"Error extracting text from MSG file {file_path}: {e}")
            # Return a basic message rather than failing completely
            return f"MSG file processing error: {Path(file_path).name} - {str(e)}"
    
    def extract_msg_metadata(self, content: bytes, file_path: str) -> Dict[str, Any]:
        """Extract essential email metadata from .msg files
        
        Returns:
            Dictionary with essential email metadata fields only
        """
        metadata = {
            "email_sender": None,
            "email_recipients_to": None,
            "email_subject": None,
            "email_date": None,
            "email_has_attachments": False,
            "email_body_preview": None
        }
        
        try:
            # Create a temporary file since extract_msg works with file paths
            with tempfile.NamedTemporaryFile(suffix='.msg', delete=False) as temp_file:
                temp_file.write(content)
                temp_file.flush()
                
                try:
                    # Open the MSG file
                    msg = extract_msg.openMsg(temp_file.name)
                    
                    # Extract sender (truncated to 200 chars)
                    if hasattr(msg, 'sender') and msg.sender:
                        metadata["email_sender"] = str(msg.sender).strip()[:200]
                    
                    # Extract recipients (truncated to 300 chars)
                    if hasattr(msg, 'to') and msg.to:
                        metadata["email_recipients_to"] = str(msg.to).strip()[:300]
                    
                    # Extract subject (truncated to 200 chars)
                    if hasattr(msg, 'subject') and msg.subject:
                        metadata["email_subject"] = str(msg.subject).strip()[:200]
                    
                    # Extract date
                    if hasattr(msg, 'date') and msg.date:
                        metadata["email_date"] = str(msg.date)
                    
                    # Check for attachments
                    if hasattr(msg, 'attachments') and msg.attachments:
                        metadata["email_has_attachments"] = True
                    
                    # Extract body preview (first 150 chars)
                    body_text = ""
                    if hasattr(msg, 'body') and msg.body:
                        body_text = str(msg.body).strip()
                    elif hasattr(msg, 'htmlBody') and msg.htmlBody:
                        # Basic HTML cleanup for preview
                        import re
                        html_text = re.sub(r'<[^>]+>', '', str(msg.htmlBody))
                        body_text = html_text.replace('&nbsp;', ' ').replace('&amp;', '&').strip()
                    
                    if body_text:
                        metadata["email_body_preview"] = body_text[:150]
                    
                    # Close the MSG file
                    msg.close()
                    
                    self.logger.info(f"Successfully extracted essential email metadata from: {Path(file_path).name}")
                    
                finally:
                    # Clean up temporary file
                    import os
                    try:
                        os.unlink(temp_file.name)
                    except:
                        pass
                        
        except Exception as e:
            self.logger.error(f"Error extracting email metadata from {file_path}: {e}")
        
        return metadata
    
    def _extract_doc_text(self, content: bytes, file_path: str) -> str:
        """Extract text from legacy .doc files using multiple fallback methods"""
        try:
            # Method 1: Try python-docx2txt if available
            try:
                import docx2txt
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    
                    text = docx2txt.process(temp_file.name)
                    
                    # Clean up
                    import os
                    os.unlink(temp_file.name)
                    
                    if text and text.strip():
                        self.logger.debug(f"Successfully extracted .doc text using docx2txt: {len(text)} chars")
                        return text.strip()
                        
            except ImportError:
                self.logger.debug("docx2txt not available, trying alternative methods...")
            except Exception as e:
                self.logger.debug(f"docx2txt extraction failed: {e}")
            
            # Method 2: Try antiword command line tool if available
            try:
                with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as temp_file:
                    temp_file.write(content)
                    temp_file.flush()
                    
                    # Try using antiword command
                    result = subprocess.run(
                        ['antiword', temp_file.name],
                        capture_output=True,
                        text=True,
                        timeout=30
                    )
                    
                    # Clean up
                    import os
                    os.unlink(temp_file.name)
                    
                    if result.returncode == 0 and result.stdout.strip():
                        self.logger.debug(f"Successfully extracted .doc text using antiword: {len(result.stdout)} chars")
                        return result.stdout.strip()
                    
            except (subprocess.TimeoutExpired, FileNotFoundError, subprocess.CalledProcessError) as e:
                self.logger.debug(f"antiword extraction failed: {e}")
            except Exception as e:
                self.logger.debug(f"antiword extraction error: {e}")
            
            # Method 3: Try textract if available (requires system dependencies)
            try:
                import textract
                text = textract.process(io.BytesIO(content), extension='.doc').decode('utf-8')
                if text and text.strip():
                    self.logger.debug(f"Successfully extracted .doc text using textract: {len(text)} chars")
                    return text.strip()
                    
            except ImportError:
                self.logger.debug("textract not available for .doc processing")
            except Exception as e:
                self.logger.debug(f"textract extraction failed: {e}")
            
            # Method 4: Basic fallback - try to extract readable text from binary
            try:
                # Simple fallback: look for readable text in the binary content
                text_content = content.decode('utf-8', errors='ignore')
                
                # Filter out control characters and keep only printable text
                import string
                printable_chars = set(string.printable)
                readable_text = ''.join(char for char in text_content if char in printable_chars)
                
                # Extract meaningful sentences (lines with reasonable length)
                lines = [line.strip() for line in readable_text.split('\n') if line.strip()]
                meaningful_lines = [line for line in lines if len(line) > 10 and not line.startswith('\\x')]
                
                if meaningful_lines:
                    extracted_text = '\n'.join(meaningful_lines[:50])  # Limit to first 50 meaningful lines
                    self.logger.warning(f"Using basic text extraction for .doc file: {len(extracted_text)} chars")
                    return extracted_text
                    
            except Exception as e:
                self.logger.debug(f"Basic text extraction failed: {e}")
            
            # If all methods fail, return a placeholder
            self.logger.warning(f"Could not extract text from .doc file: {Path(file_path).name}")
            return f"[Legacy Word Document: {Path(file_path).name}] - Text extraction not available. Install docx2txt or antiword for better .doc support."
            
        except Exception as e:
            self.logger.error(f"Error processing .doc file {file_path}: {e}")
            return f"[Error processing .doc file: {Path(file_path).name}]"
    
    def _extract_pptx_text(self, content: bytes, file_path: str) -> str:
        """Extract text from PowerPoint (.pptx) files with optional VLM enhancement"""
        try:
            # Use enhanced parser if available (with VLM support)
            if self.enhanced_powerpoint_parser and self.enhanced_powerpoint_parser.can_process(file_path):
                self.logger.info(f"🔍 Using enhanced PowerPoint parser with VLM for: {Path(file_path).name}")
                result = self.enhanced_powerpoint_parser.extract_enhanced_content(content, file_path)
                text_content = result.get("text_content", "")
                
                # Log VLM analysis results
                if result.get("total_images", 0) > 0 or result.get("total_charts", 0) > 0:
                    self.logger.info(f"📊 VLM analysis: {result.get('total_images', 0)} images, "
                                   f"{result.get('total_charts', 0)} charts analyzed")
                
                return text_content if text_content.strip() else f"Enhanced PowerPoint analysis: {Path(file_path).name} (no extractable content)"
            
            # Fall back to standard parser
            elif self.powerpoint_parser.can_process(file_path):
                self.logger.debug(f"📄 Using standard PowerPoint parser for: {Path(file_path).name}")
                text_content = self.powerpoint_parser.extract_text(content, file_path)
                return text_content if text_content.strip() else f"PowerPoint presentation: {Path(file_path).name} (no extractable text content)"
            
            else:
                return "PowerPoint parsing not available - missing python-pptx dependency"
            
        except Exception as e:
            self.logger.error(f"Error extracting PowerPoint content from {file_path}: {e}")
            return f"Error extracting PowerPoint content: {str(e)}"
    
    def get_file_info(self, file_path: str) -> dict:
        """Get information about file processing capabilities"""
        extension = Path(file_path).suffix.lower()
        
        info = {
            "extension": extension,
            "can_process": self.can_process(file_path),
            "processing_method": self.supported_extensions.get(extension, "unsupported"),
            "output_type": "pdf" if extension in ['.pdf', '.png', '.jpg', '.jpeg'] else "text"
        }
        
        return info 